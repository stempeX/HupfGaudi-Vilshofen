"""
Prüfprotokoll-Generator für aufblasbare Spielgeräte
gemäß EN 14960-1:2019 | BetrSichV | TRBS 1203
Erzeugt: Pruefprotokoll_Huepfburg.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === FARBEN ===
ORANGE = RGBColor(0xFF, 0x5A, 0x1F)
DARK   = RGBColor(0x2C, 0x2C, 0x2C)
GRAY   = RGBColor(0x6C, 0x6C, 0x6C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xCC, 0x00, 0x00)
GREEN  = RGBColor(0x4C, 0xAF, 0x50)
AMBER  = RGBColor(0xFF, 0x98, 0x00)

# === DOKUMENT SETUP ===
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)

# === HILFSFUNKTIONEN ===
def spacing(para, b=0, a=40):
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(b))
    sp.set(qn('w:after'), str(a))
    para._p.get_or_add_pPr().append(sp)

def shading(para, fill):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    para._p.get_or_add_pPr().append(s)

def cell_bg(cell, fill):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(s)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_cm * 567)))
    w.set(qn('w:type'), 'dxa')
    tcPr.append(w)

def full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), '5000')
    w.set(qn('w:type'), 'pct')
    tblPr.append(w)

def add_run(para, text, bold=False, size=9.5, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.font.color.rgb = color or DARK
    return r

def para(doc, text, bold=False, size=9.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, b=30, a=30):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, b, a)
    add_run(p, text, bold, size, color)
    return p

def section_header(doc, text, bg='FF5A1F'):
    p = doc.add_paragraph()
    spacing(p, 120, 80)
    shading(p, bg)
    r = p.add_run('  ' + text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.color.rgb = WHITE

def orange_line(doc):
    p = doc.add_paragraph()
    spacing(p, 40, 40)
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), 'FF5A1F')
    bdr.append(bot)
    p._p.get_or_add_pPr().append(bdr)

def bottom_border_para(para, color='AAAAAA'):
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:color'), color)
    bdr.append(bot)
    para._p.get_or_add_pPr().append(bdr)

def form_field(doc, label):
    p_l = doc.add_paragraph()
    spacing(p_l, 30, 2)
    add_run(p_l, label, True, 8, GRAY)
    p_f = doc.add_paragraph()
    spacing(p_f, 0, 8)
    bottom_border_para(p_f)

def form_field_inline(doc, labels):
    """Zwei Felder nebeneinander als Tabelle"""
    t = doc.add_table(rows=len(labels), cols=4)
    full_width(t)
    for i, (l1, l2) in enumerate(labels):
        r = t.rows[i]
        p1 = r.cells[0].paragraphs[0]
        add_run(p1, l1, True, 8, GRAY)
        bottom_border_para(r.cells[1].paragraphs[0])
        p2 = r.cells[2].paragraphs[0]
        add_run(p2, '  ' + l2, True, 8, GRAY)
        bottom_border_para(r.cells[3].paragraphs[0])

def checkbox_line(doc, text):
    p = doc.add_paragraph()
    spacing(p, 20, 20)
    add_run(p, '\u2610  ', False, 10)
    add_run(p, text, False, 9.5)

def ja_nein_line(doc, text):
    p = doc.add_paragraph()
    spacing(p, 20, 20)
    add_run(p, '\u2610 Ja   \u2610 Nein   ', False, 9.5)
    add_run(p, text, False, 9.5)


# =============================================
# SEITE 1: TITEL
# =============================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_title, 0, 60)
shading(p_title, 'FF5A1F')
add_run(p_title, 'HUPFGAUDI VILSHOFEN', True, 18, WHITE)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_sub, 20, 20)
add_run(p_sub, 'PRÜFPROTOKOLL AUFBLASBARE SPIELGERÄTE', True, 14, ORANGE)

p_norm = doc.add_paragraph()
p_norm.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_norm, 0, 40)
add_run(p_norm, 'gemäß EN 14960-1:2019  |  BetrSichV  |  TRBS 1203', False, 9, GRAY)

p_addr = doc.add_paragraph()
p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_addr, 20, 60)
add_run(p_addr, 'Christoph Gutsmidl  \u2022  Böcklbacher Str. 7  \u2022  94474 Vilshofen (Alkofen)', False, 9, GRAY)

orange_line(doc)

# Protokoll-Nr. und Datum
form_field(doc, 'Protokoll-Nr.:')
form_field(doc, 'Prüfdatum:')

para(doc, 'Prüfart:', True, 9.5, GRAY, b=40, a=10)
p = doc.add_paragraph()
spacing(p, 0, 40)
add_run(p, '\u2610 Jahreshauptinspektion     \u2610 Sichtprüfung vor Inbetriebnahme     \u2610 Sonderprüfung', False, 9.5)

# =============================================
# ABSCHNITT 1: PRÜFER
# =============================================
section_header(doc, '1. PRÜFER / BEFÄHIGTE PERSON')
form_field(doc, 'Name:')
form_field(doc, 'Firma:')
form_field(doc, 'Qualifikation:')
form_field(doc, 'Telefon / E-Mail:')
form_field(doc, 'Grundlage der Befähigung (Ausbildung, Erfahrung, Norm-Kenntnisse):')

# =============================================
# ABSCHNITT 2: AUFTRAGGEBER
# =============================================
section_header(doc, '2. AUFTRAGGEBER / BETREIBER')
form_field(doc, 'Name / Firma:')
form_field(doc, 'Ansprechpartner:')
form_field(doc, 'Adresse:')
form_field(doc, 'Telefon / E-Mail:')
form_field(doc, 'Prüfort / Standort:')

# =============================================
# ABSCHNITT 3: GERÄTEIDENTIFIKATION
# =============================================
section_header(doc, '3. GERÄTEIDENTIFIKATION')
form_field(doc, 'Bezeichnung / Typ:')
form_field(doc, 'Hersteller:')
form_field(doc, 'Seriennummer Gerät:')
form_field(doc, 'Seriennummer Gebläse:')
form_field(doc, 'Baujahr / Anschaffungsdatum:')
form_field(doc, 'Abmessungen (L × B × H):')
form_field(doc, 'Gewicht:')
form_field(doc, 'Max. Nutzerzahl / Altersgruppe:')

p = doc.add_paragraph()
spacing(p, 40, 20)
add_run(p, '\u2610 CE-Kennzeichnung vorhanden     ', False, 9.5)
add_run(p, '\u2610 EN 14960-Zertifikat vorhanden     ', False, 9.5)
add_run(p, '\u2610 Betriebsanleitung vorhanden', False, 9.5)

# =============================================
# ABSCHNITT 4: DOKUMENTENPRÜFUNG
# =============================================
section_header(doc, '4. DOKUMENTENPRÜFUNG')
ja_nein_line(doc, 'Betriebsanleitung des Herstellers vorhanden und vollständig')
ja_nein_line(doc, 'Vorherige Prüfberichte vorhanden')
ja_nein_line(doc, 'Herstellerzertifikate / Konformitätserklärung vorhanden')
ja_nein_line(doc, 'Reparatur- / Wartungshistorie dokumentiert')

# =============================================
# ABSCHNITT 5: PRÜFPUNKTE
# =============================================
section_header(doc, '5. PRÜFPUNKTE NACH EN 14960-1:2019')

p_info = doc.add_paragraph()
spacing(p_info, 40, 60)
shading(p_info, 'FFF8E7')
add_run(p_info, '  Bewertung: ', True, 9)
add_run(p_info, 'i.O. = in Ordnung  |  n.i.O. = nicht in Ordnung  |  n.a. = nicht anwendbar', False, 9)

PRUEFPUNKTE = [
    ('1',  'Material / Gewebe',       'Verschleiß, Risse, Löcher, UV-Schäden, Verfärbungen, Schimmelbefall'),
    ('2',  'Nähte',                    'Bodennähte, Wandnähte, Turmnähte – Auflösung, Beschädigung, Fadenbruch'),
    ('3',  'Innere Verbindungen',      'Verstrebungen, Kammern, Zwischenwände – Zustand und Befestigung'),
    ('4',  'Wände und Türme',          'Sichere Befestigung, Geradheit, keine Verformung oder Neigung'),
    ('5',  'Verankerungssystem',       'Ösen, Gurte, Seile – Verschleiß, Risse, Abrieb, Korrosion'),
    ('6',  'Erdanker / Ballast',       'Anzahl und Art normkonform, ausreichende Haltekraft'),
    ('7',  'Luftdruck',                'Betriebsdruck 10–20 mbar, stabile Standfestigkeit, keine Druckverluste'),
    ('8',  'Gebläse',                  'Funktion, Zustand, Leistung, Geräuschpegel, sichere Befestigung'),
    ('9',  'Schutzgitter Gebläse',     'Ein- und Austrittsöffnung geschützt, keine Zugangsmöglichkeit'),
    ('10', 'Elektrische Leitungen',    'Kabel, Stecker, Isolation – Zustand, Beschädigung, Schutzleiter'),
    ('11', 'Fangstellen',              'Keine Kopf-, Finger-, Hals- oder Kleidungsfangstellen'),
    ('12', 'Ein- / Ausstieg',          'Sicher, ausreichend dimensioniert, keine Stolperstellen'),
    ('13', 'Kennzeichnung',            'Typenschild, max. Nutzerzahl, Altersangabe, Herstellerdaten lesbar'),
    ('14', 'Sicherheitsabstände',      'Aufstellfläche ausreichend, Mindestabstände eingehalten'),
]

table = doc.add_table(rows=len(PRUEFPUNKTE) + 1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
full_width(table)

# Header
headers = ['Nr.', 'Prüfpunkt', 'i.O.', 'n.i.O.', 'n.a.', 'Bemerkungen']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell_bg(cell, 'FF5A1F')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.name = 'Calibri'
    r.font.color.rgb = WHITE

# Spaltenbreiten
for row in table.rows:
    set_cell_width(row.cells[0], 1.0)   # Nr
    set_cell_width(row.cells[1], 6.5)   # Prüfpunkt
    set_cell_width(row.cells[2], 1.2)   # i.O.
    set_cell_width(row.cells[3], 1.2)   # n.i.O.
    set_cell_width(row.cells[4], 1.2)   # n.a.
    set_cell_width(row.cells[5], 5.5)   # Bemerkungen

# Datenzeilen
for idx, (nr, titel, beschreibung) in enumerate(PRUEFPUNKTE):
    row = table.rows[idx + 1]
    if idx % 2 == 1:
        for cell in row.cells:
            cell_bg(cell, 'FDF6F0')

    # Nr.
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, nr, True, 9)

    # Prüfpunkt
    p = row.cells[1].paragraphs[0]
    add_run(p, titel, True, 9)
    add_run(p, '\n' + beschreibung, False, 8, GRAY)

    # Checkboxen
    for ci in (2, 3, 4):
        p = row.cells[ci].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, '\u2610', False, 12)

    # Bemerkungen leer lassen

# Tabellenrahmen
tbl = table._tbl
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
borders = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    e = OxmlElement(f'w:{side}')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:color'), 'DDDDDD')
    borders.append(e)
tblPr.append(borders)

# =============================================
# ABSCHNITT 6: MÄNGELBESCHREIBUNG
# =============================================
section_header(doc, '6. FESTGESTELLTE MÄNGEL', 'CC0000')

mt = doc.add_table(rows=6, cols=5)
full_width(mt)
mheaders = ['Nr.', 'Beschreibung des Mangels', 'Bewertung', 'Empfohlene Maßnahme', 'Frist']
for i, h in enumerate(mheaders):
    cell = mt.rows[0].cells[i]
    cell_bg(cell, 'CC0000')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.name = 'Calibri'
    r.font.color.rgb = WHITE

for row_idx in range(1, 6):
    p = mt.rows[row_idx].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(row_idx), False, 9)

# Rahmen Mängeltabelle
tbl2 = mt._tbl
tblPr2 = tbl2.find(qn('w:tblPr'))
if tblPr2 is None:
    tblPr2 = OxmlElement('w:tblPr')
    tbl2.insert(0, tblPr2)
borders2 = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    e = OxmlElement(f'w:{side}')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:color'), 'DDDDDD')
    borders2.append(e)
tblPr2.append(borders2)

# Mängelbewertung Info
p_mb = doc.add_paragraph()
spacing(p_mb, 60, 40)
shading(p_mb, 'FFF8E7')
add_run(p_mb, '  Mängelbewertung: ', True, 9)
add_run(p_mb, 'Gering', True, 9, GREEN)
add_run(p_mb, ' = Betrieb möglich  |  ', False, 9)
add_run(p_mb, 'Erheblich', True, 9, AMBER)
add_run(p_mb, ' = zeitnahe Behebung  |  ', False, 9)
add_run(p_mb, 'Gefährlich', True, 9, RED)
add_run(p_mb, ' = sofortige Stilllegung', False, 9)

# =============================================
# ABSCHNITT 7: GESAMTBEWERTUNG
# =============================================
section_header(doc, '7. GESAMTBEWERTUNG')

p = doc.add_paragraph()
spacing(p, 60, 30)
add_run(p, '\u2610  ', False, 12, GREEN)
add_run(p, 'Keine Mängel', True, 10, GREEN)
add_run(p, ' – Das Gerät ist betriebssicher und kann uneingeschränkt genutzt werden.', False, 9.5)

p = doc.add_paragraph()
spacing(p, 30, 30)
add_run(p, '\u2610  ', False, 12, RGBColor(0x8B, 0xC3, 0x4A))
add_run(p, 'Geringe Mängel', True, 10, RGBColor(0x8B, 0xC3, 0x4A))
add_run(p, ' – Betrieb unter Auflagen möglich. Mängel bei nächster Wartung beheben.', False, 9.5)

p = doc.add_paragraph()
spacing(p, 30, 30)
add_run(p, '\u2610  ', False, 12, AMBER)
add_run(p, 'Erhebliche Mängel', True, 10, AMBER)
add_run(p, ' – Betrieb eingeschränkt. Nachprüfung nach Mängelbehebung erforderlich.', False, 9.5)

p = doc.add_paragraph()
spacing(p, 30, 40)
add_run(p, '\u2610  ', False, 12, RED)
add_run(p, 'Gefährliche Mängel', True, 10, RED)
add_run(p, ' – Sofortige Stilllegung! Betrieb bis Mängelbehebung und Nachprüfung nicht zulässig.', False, 9.5)

# =============================================
# ABSCHNITT 8: NÄCHSTE PRÜFUNG
# =============================================
section_header(doc, '8. NÄCHSTE PRÜFUNG')
form_field(doc, 'Empfohlener Termin:')
form_field(doc, 'Prüfart:')
form_field(doc, 'Besondere Hinweise:')

# =============================================
# ABSCHNITT 9: UNTERSCHRIFTEN
# =============================================
section_header(doc, '9. UNTERSCHRIFTEN')

sig = doc.add_table(rows=3, cols=2)
full_width(sig)
# Leerzeile für Unterschrift
for cell in sig.rows[0].cells:
    p = cell.paragraphs[0]
    spacing(p, 200, 0)

# Linie
for cell in sig.rows[1].cells:
    p = cell.paragraphs[0]
    bottom_border_para(p, '333333')

# Labels
p1 = sig.rows[2].cells[0].paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p1, 'Prüfer / Befähigte Person', True, 8, GRAY)
r1 = p1.add_run('\nOrt, Datum')
r1.font.size = Pt(7.5)
r1.font.name = 'Calibri'
r1.font.color.rgb = GRAY

p2 = sig.rows[2].cells[1].paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, 'Auftraggeber / Betreiber', True, 8, GRAY)
r2 = p2.add_run('\nOrt, Datum')
r2.font.size = Pt(7.5)
r2.font.name = 'Calibri'
r2.font.color.rgb = GRAY

# Rahmen entfernen bei Unterschriftentabelle
tbl3 = sig._tbl
tblPr3 = tbl3.find(qn('w:tblPr'))
if tblPr3 is None:
    tblPr3 = OxmlElement('w:tblPr')
    tbl3.insert(0, tblPr3)
borders3 = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    e = OxmlElement(f'w:{side}')
    e.set(qn('w:val'), 'none')
    borders3.append(e)
tblPr3.append(borders3)

# =============================================
# FUSSZEILE
# =============================================
orange_line(doc)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_foot, 40, 20)
add_run(p_foot, 'HupfGaudi Vilshofen', True, 8, ORANGE)
add_run(p_foot, '  |  Christoph Gutsmidl  |  Böcklbacher Str. 7, 94474 Vilshofen  |  +49 151 28861367  |  hupfgaudi@gmail.com', False, 7.5, GRAY)

p_foot2 = doc.add_paragraph()
p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_foot2, 0, 0)
add_run(p_foot2, 'Prüfung gemäß EN 14960-1:2019  |  Befähigte Person nach BetrSichV / TRBS 1203', False, 7.5, GRAY)

# =============================================
# SPEICHERN
# =============================================
output = 'Pruefprotokoll_Huepfburg.docx'
doc.save(output)
print(f'Prüfprotokoll erstellt: {output}')
