from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xFF,0x6B,0x00)
DARK   = RGBColor(0x2C,0x2C,0x2C)
GRAY   = RGBColor(0x6C,0x6C,0x6C)
WHITE  = RGBColor(0xFF,0xFF,0xFF)
RED    = RGBColor(0xCC,0x00,0x00)
BLUE   = RGBColor(0x33,0x66,0x99)

doc = Document()
sec = doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7)
sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)

def spacing(para, b=0, a=40):
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(b))
    sp.set(qn('w:after'), str(a))
    para._p.get_or_add_pPr().append(sp)

def shading(para, fill):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    para._p.get_or_add_pPr().append(s)

def cell_bg(cell, fill):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(s)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    b = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'none')
        b.append(e)
    tblPr.append(b)

def full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), '9360')
    w.set(qn('w:type'), 'dxa')
    tblPr.append(w)

def orange_line(doc):
    p = doc.add_paragraph()
    spacing(p, 60, 60)
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), 'FF6B00')
    bdr.append(bot)
    p._p.get_or_add_pPr().append(bdr)

def section_header(doc, text, bg='FF6B00'):
    p = doc.add_paragraph()
    spacing(p, 100, 80)
    shading(p, bg)
    r = p.add_run('  ' + text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.color.rgb = WHITE

def add_run(para, text, bold=False, size=9.5, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.font.color.rgb = color or DARK
    return r

def para(doc, text, bold=False, size=9.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, b=30, a=30):
    p = doc.add_paragraph()
    p.alignment = align
    spacing(p, b, a)
    add_run(p, text, bold, size, color)
    return p

def bullet(doc, prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    spacing(p, 20, 20)
    if prefix:
        add_run(p, prefix, True, 9.5)
    add_run(p, text, False, 9.5)

def bottom_border_para(para, color='AAAAAA'):
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:color'), color)
    bdr.append(bot)
    para._p.get_or_add_pPr().append(bdr)

def top_border_para(para, color='FF6B00'):
    bdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')
    top.set(qn('w:color'), color)
    bdr.append(top)
    para._p.get_or_add_pPr().append(bdr)

def form_fields(doc, labels):
    for label in labels:
        p_l = doc.add_paragraph()
        spacing(p_l, 30, 2)
        add_run(p_l, label, True, 8, GRAY)
        p_f = doc.add_paragraph()
        spacing(p_f, 0, 8)
        bottom_border_para(p_f)

# ===== SEITE 1: Titel & Vertragsparteien =====
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_title, 0, 120)
shading(p_title, 'FF6B00')
add_run(p_title, 'HUPFGAUDI VILSHOFEN', True, 20, WHITE)
add_run(p_title, '   |   MIETVERTRAG', False, 14, RGBColor(0xFF,0xE0,0xC0))

p_addr = doc.add_paragraph()
p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p_addr, 40, 100)
add_run(p_addr, 'Christoph & Yasmin Gutsmidl  \u2022  Böcklbacher Str. 7  \u2022  94474 Vilshofen (Alkofen)', False, 9, GRAY)

orange_line(doc)

section_header(doc, 'VERMIETER')
p = doc.add_paragraph()
spacing(p, 60, 60)
add_run(p, 'Christoph & Yasmin Gutsmidl\nBöcklbacher Str. 7\n94474 Vilshofen (Alkofen)', True, 10)

section_header(doc, 'MIETER UND RECHNUNGSADRESSE')
para(doc, '(nachstehend Mieter genannt, ausgewiesen durch Personalausweis)', size=9, color=GRAY, b=20, a=20)
form_fields(doc, ['Name, Vorname','Firma','Straße + Hausnummer','PLZ + Ort','Telefon / Mobil','E-Mail'])

section_header(doc, 'VERANSTALTUNGSADRESSE  (falls abweichend von der Rechnungsadresse)')
form_fields(doc, ['Name, Vorname','Firma','Straße + Hausnummer','PLZ + Ort','Telefon / Mobil','E-Mail'])

doc.add_page_break()

# ===== SEITE 2: Mietgegenstand =====
section_header(doc, 'MIETGEGENSTAND & KONDITIONEN')
para(doc, 'Die Beschaffung von eventuell erforderlichen Genehmigungen (z. B. Aufstellen auf öffentlichen Flächen) obliegt allein im Verantwortungsbereich des Mieters.', size=9, color=GRAY, b=60, a=60)

t_hb = doc.add_table(rows=2, cols=3)
full_width(t_hb)
for c, h in zip(t_hb.rows[0].cells, ['Hüpfburg / Modell','Beschreibung','Mietpreis (€)']):
    cell_bg(c, 'FFE0C0')
    add_run(c.paragraphs[0], h, True, 9.5, ORANGE)
for c in t_hb.rows[1].cells:
    c.paragraphs[0].paragraph_format.space_before = Pt(10)
    c.paragraphs[0].paragraph_format.space_after = Pt(10)

p_ew = doc.add_paragraph()
spacing(p_ew, 80, 20)
add_run(p_ew, 'Einsatzzweck  ', True, 9.5, ORANGE)
add_run(p_ew, '(gewerbliche Nutzung ausgeschlossen)', False, 9, GRAY)
p_ew2 = doc.add_paragraph()
spacing(p_ew2, 10, 40)
for o in ['\u2610 Kindergeburtstag','\u2610 Hochzeit','\u2610 Betriebsfeier','\u2610 Vereinsfeier','\u2610 Sonstiges: ___________']:
    add_run(p_ew2, o+'   ', False, 10)

p_ug = doc.add_paragraph()
spacing(p_ug, 40, 20)
add_run(p_ug, 'Aufstelluntergrund  ', True, 9.5, ORANGE)
p_ug2 = doc.add_paragraph()
spacing(p_ug2, 10, 40)
for o in ['\u2610 Rasen','\u2610 Zement','\u2610 Asphalt','\u2610 Sand','\u2610 Kies','\u2610 Pflaster','\u2610 Sonstiges: _______']:
    add_run(p_ug2, o+'   ', False, 10)

orange_line(doc)

p_sa = doc.add_paragraph()
spacing(p_sa, 60, 20)
add_run(p_sa, '\u2610  SELBSTABHOLUNG', True, 10)
t_sa = doc.add_table(rows=2, cols=2)
full_width(t_sa)
for c, h in zip(t_sa.rows[0].cells, ['Abholung \u2013 Datum + Uhrzeit','Rückgabe \u2013 Datum + Uhrzeit']):
    cell_bg(c, 'F5F5F5')
    add_run(c.paragraphs[0], h, True, 8, GRAY)
for c in t_sa.rows[1].cells:
    c.paragraphs[0].paragraph_format.space_before = Pt(10)
    c.paragraphs[0].paragraph_format.space_after = Pt(10)
para(doc, 'Die Abholung, der Auf- und Abbau sowie der Rücktransport erfolgt durch den Mieter. Ein Anhänger oder Großraum-PKW ist notwendig. Die Mietsache muss trocken und gereinigt zurückgegeben werden.', size=9, color=GRAY, b=40, a=40)

p_ls = doc.add_paragraph()
spacing(p_ls, 60, 20)
add_run(p_ls, '\u2610  LIEFER- & AUFBAUSERVICE  +  ', True, 10)
add_run(p_ls, '_________ \u20ac  pauschal je Einsatz', True, 10, ORANGE)
t_ls = doc.add_table(rows=2, cols=2)
full_width(t_ls)
for c, h in zip(t_ls.rows[0].cells, ['Lieferung \u2013 Datum + Uhrzeit','Abholung \u2013 Datum + Uhrzeit']):
    cell_bg(c, 'F5F5F5')
    add_run(c.paragraphs[0], h, True, 8, GRAY)
for c in t_ls.rows[1].cells:
    c.paragraphs[0].paragraph_format.space_before = Pt(10)
    c.paragraphs[0].paragraph_format.space_after = Pt(10)

orange_line(doc)
p_k = doc.add_paragraph()
spacing(p_k, 60, 20)
shading(p_k, 'FFF3E0')
add_run(p_k, '  KAUTION BEI ABHOLUNG: ', True, 10, ORANGE)
add_run(p_k, '150,\u2013 \u20ac', True, 14, RED)

p_k2 = doc.add_paragraph()
spacing(p_k2, 20, 10)
shading(p_k2, 'FFF3E0')
add_run(p_k2, '  Die Kaution wird vollständig zurückerstattet, wenn:', True, 10)

for bedingung in ['die Hüpfburg sauber zurückgegeben wird', 'sie ausgekehrt und gereinigt wurde', 'ein kurzer Videobeweis vorliegt']:
    p_bed = doc.add_paragraph()
    spacing(p_bed, 6, 6)
    shading(p_bed, 'FFF3E0')
    add_run(p_bed, '     \u2713  ' + bedingung, False, 9.5)

doc.add_page_break()

# ===== SEITE 3: Kaution & Unterschrift =====
section_header(doc, 'KAUTIONSREGELUNG & VERTRAGSABSCHLUSS')
para(doc, 'Ein Kautionseinbehalt kann vorgenommen werden bei:', True, b=60, a=20)
for x in ['stark verschmutzter oder nasser Mietsache', 'zu spät erfolgter Rückgabe', 'Beschädigungen an der Mietsache', 'fehlendem Videobeweis der Reinigung']:
    bullet(doc, '', x)

para(doc, 'Die Kaution (150,\u2013 \u20ac) wird vollständig zurückerstattet, wenn die Hüpfburg sauber zurückgegeben, ausgekehrt und gereinigt wurde und ein kurzer Videobeweis vorliegt. Die Rückerstattung erfolgt bar bei Übergabe oder bis spätestens eine Woche nach Rückgabe.', b=60, a=60)

p_gk = doc.add_paragraph()
spacing(p_gk, 60, 80)
shading(p_gk, 'FF6B00')
add_run(p_gk, '  GESAMTMIETKOSTEN: ', True, 11, WHITE)
add_run(p_gk, '________________________________ \u20ac', True, 14, RGBColor(0xFF,0xE0,0xA0))
add_run(p_gk, '   (inkl. MwSt.)', False, 9, RGBColor(0xFF,0xE0,0xA0))

para(doc, 'Der Mieter ist einverstanden mit den Inhalten dieses Vertrages, den Allgemeinen Geschäftsbedingungen, den Nutzungs- und Sicherheitsbestimmungen, der Haftungsübernahmeerklärung und hat die Widerrufsbelehrung zur Kenntnis genommen sowie Kopien aller Unterlagen erhalten.', b=60, a=60)

p_foto = doc.add_paragraph()
spacing(p_foto, 40, 60)
shading(p_foto, 'FFF3E0')
add_run(p_foto, '  \u2610  Einverständnis zur Verwendung von Fotos als Referenz (anonymisiert, ohne Namen/Standort, jederzeit widerrufbar)', False, 9.5)

orange_line(doc)
t_sig = doc.add_table(rows=2, cols=2)
no_borders(t_sig)
full_width(t_sig)
for c, lbl in zip(t_sig.rows[0].cells, ['Datum','Datum']):
    add_run(c.paragraphs[0], lbl, True, 8, GRAY)
    c.add_paragraph('___ . ___ . _________').runs[0].font.size = Pt(10)
for c, lbl in zip(t_sig.rows[1].cells, ['Unterschrift Mieter','Unterschrift Vermieter']):
    top_border_para(c.paragraphs[0])
    add_run(c.paragraphs[0], lbl, True, 9.5, ORANGE)

doc.add_page_break()

# ===== SEITE 4-5: Haftungsübernahmeerklärung =====
section_header(doc, 'HAFTUNGSÜBERNAHMEERKLÄRUNG')
items = [
    ('Übernahmezustand: ', 'Der Mieter hat die Mietsachen im einwandfreiem und gebrauchsfähigem Zustand übernommen. Eventuelle Schäden sind in der Mängelliste aufzuführen. Später vorgebrachte Einwendungen können nicht anerkannt werden.'),
    ('Sicherheitsvorschriften: ', 'Den Nutzungs- und Sicherheitsbestimmungen sowie Auf- und Abbauanweisungen des Vermieters ist Folge zu leisten.'),
    ('Sorgfaltspflicht: ', 'Der Mieter verpflichtet sich, mit der Mietsache pfleglich sowie sorgfältig umzugehen, keine Veränderungen vorzunehmen und sauber zurückzugeben.'),
    ('Eigenes Risiko: ', 'Die Nutzung der Mietsache erfolgt ausschließlich auf eigene Gefahr.'),
    ('Unfallmeldung: ', 'Bei einem Unfall hat der Mieter den Vermieter sogleich schriftlich zu unterrichten. Gegnerische Ansprüche dürfen nicht anerkannt werden.'),
    ('Haftungsumfang: ', 'Der Mieter trägt die volle Verantwortung für die Verkehrssicherungspflicht und haftet für Verschmutzung, Diebstahl, Sach- und Personenschäden, Fehlbedienung und Vandalismus.'),
    ('Versicherung: ', 'Der Mieter ist für angemessenen Versicherungsschutz selbst verantwortlich (privat: Haftpflicht; Betrieb: Betriebshaftpflicht).'),
    ('Freistellung: ', 'Der Mieter entbindet den Vermieter von jeglichen Kosten und Rechtsanwaltskosten durch Klagen Dritter.'),
]
for pre, txt in items:
    bullet(doc, pre, txt)

orange_line(doc)
section_header(doc, 'SCHADENSPAUSCHALEN', 'CC0000')
t_sp = doc.add_table(rows=5, cols=2)
full_width(t_sp)
sp_data = [
    ('Art des Schadens', 'Pauschale'),
    ('Kleine Beschädigung (Risse max. 5x5 cm, nicht sicherheitsrelevant)', '150,\u2013 \u20ac'),
    ('Hüpfburg \u2013 irreparabler Schaden', '1.500,\u2013 \u20ac'),
    ('Profi-Hüpfburg / Eventmodule \u2013 irreparabel', 'bis 8.000,\u2013 \u20ac'),
    ('Maschinen / Gebläse', '300,\u2013 \u20ac'),
]
for i, (c1, c2) in enumerate(sp_data):
    for c, txt in zip(t_sp.rows[i].cells, [c1, c2]):
        c.text = txt
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        c.paragraphs[0].runs[0].font.name = 'Calibri'
        if i == 0:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            cell_bg(c, 'CC0000')
        else:
            c.paragraphs[0].runs[0].font.color.rgb = DARK
            cell_bg(c, 'FFF5F5' if i % 2 == 0 else 'FFFFFF')
    if i > 0:
        t_sp.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        t_sp.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RED

para(doc, '', b=40, a=20)
t_sh = doc.add_table(rows=2, cols=2)
no_borders(t_sh)
full_width(t_sh)
add_run(t_sh.rows[0].cells[0].paragraphs[0], 'Datum', True, 8, GRAY)
t_sh.rows[0].cells[0].add_paragraph('___ . ___ . _________').runs[0].font.size = Pt(10)
top_border_para(t_sh.rows[1].cells[0].paragraphs[0], 'CC0000')
add_run(t_sh.rows[1].cells[0].paragraphs[0], 'Unterschrift Mieter', True, 9.5, RED)

doc.add_page_break()

# ===== SEITE 6-8: Sicherheitsbestimmungen =====
section_header(doc, 'NUTZUNGS- UND SICHERHEITSBESTIMMUNGEN')
para(doc, 'Folgende Bestimmungen müssen eingehalten werden. Die Bereitstellung von Strom (230 V, 16 A) oder Wasseranschlüssen ist Sache des Mieters.', b=60, a=60)

for hdr, items in [
    ('Elektrisches Gebläse', [
        'Nur die Aufsichtsperson darf Zugriff zum Gebläse haben.',
        'Kabeltrommeln müssen vor Inbetriebnahme vollständig abgerollt sein.',
        'Nur feuchtigkeitsgeschütztes Außenverlängerungskabel verwenden.',
        'Das Gebläse darf nicht ohne Anschluss an die Hüpfburg eingeschaltet werden.',
        'Lufteintritt darf nicht behindert werden \u2013 keine Fremdteile eingesaugt werden.',
        'Das Gebläse muss sauber und trocken aufgestellt sein.',
        'Vor dem Aufblasen: Luftkanal im 90°-Winkel ausrichten, nicht geknickt.',
        'Keine Kinder im Bereich des Gebläses.',
    ]),
    ('Aufstellfläche', [
        'Vorzugsweise freie Gras- oder Rasenfläche wählen.',
        'Immer Schutzplane unterlegen, Fläche frei von Steinen und spitzen Gegenständen.',
        'Keine Gefahrenquellen auf der offenen Seite der Hüpfburg.',
        'Eingangs-Teppich auslegen.',
        'Nicht neben einem Swimmingpool aufstellen.',
        'Mindestens 1,8 m freier Platz rund um die Hüpfburg.',
    ]),
    ('Vorbereitung & Aufblasen', [
        'Aufbau bei starkem Wind oder Niederschlag ist zu unterlassen.',
        'Niemand darf die Hüpfburg betreten, bevor sie vollständig aufgeblasen ist.',
        'Aufsichtsperson beobachtet den gesamten Füllvorgang.',
        'Gebläse darf während der gesamten Nutzungszeit nicht abgeschaltet werden.',
    ]),
    ('Hüpfen & Verhalten', [
        'Erwachsene dürfen die Hüpfburg wegen hoher Punktbelastung NICHT benutzen.',
        'Auf vergleichbares Alter und Größe der gleichzeitig hüpfenden Kinder achten.',
        'Verboten: Speisen, Getränke, Schuhe, Lutscher, Kaugummi.',
        'Verboten: Halsketten, Ringe, Brillen, Gürtelschnallen.',
        'Verboten: Saltos, Handstände, Wrestling, Klettern an den Wänden.',
        'Verboten: Tiere in der Hüpfburg.',
        'Bei Regen, Sturm oder Gewitter: sofort Nutzung einstellen und abbauen.',
        'Bei Stromausfall: Hüpfburg sofort verlassen.',
        'Während des Luftablassens: keine Person in der Hüpfburg.',
    ]),
]:
    section_header(doc, hdr, '336699')
    for item in items:
        bullet(doc, '', item)

orange_line(doc)
para(doc, 'Die Sicherheitsbestimmungen habe ich gelesen, akzeptiert und sichere zu, diese zu beachten und einzuhalten.', b=60, a=60)
para(doc, 'Wir wünschen Ihnen und Ihren Kindern viel Spaß! \u2665', True, 10, ORANGE, WD_ALIGN_PARAGRAPH.CENTER, b=20, a=60)
t_sic = doc.add_table(rows=2, cols=2)
no_borders(t_sic)
full_width(t_sic)
add_run(t_sic.rows[0].cells[0].paragraphs[0], 'Datum', True, 8, GRAY)
t_sic.rows[0].cells[0].add_paragraph('___ . ___ . _________').runs[0].font.size = Pt(10)
top_border_para(t_sic.rows[1].cells[0].paragraphs[0], '336699')
add_run(t_sic.rows[1].cells[0].paragraphs[0], 'Unterschrift Mieter', True, 9.5, BLUE)

doc.add_page_break()

# ===== AGB =====
section_header(doc, 'ALLGEMEINE GESCHÄFTSBEDINGUNGEN (AGB)')
agb = [
    ('\u00a71 Geltungsbereich', 'Diese AGB der HupfGaudi Vilshofen gelten für alle Mietverträge. Eigene Bedingungen des Mieters werden nicht anerkannt, sofern nichts anderes vereinbart ist.'),
    ('\u00a72 Vertragsschluss', 'Angebote können online, telefonisch oder per E-Mail abgegeben werden. Annahmefrist: 5 Tage. Mündliche Absprachen haben keine Gültigkeit.'),
    ('\u00a73 Widerrufsrecht', 'Verbrauchern steht grundsätzlich ein Widerrufsrecht zu. Details in der separaten Widerrufsbelehrung.'),
    ('\u00a74 Überlassung', 'Die Mietsache bleibt Eigentum des Vermieters. Wird der Mietartikel nicht innerhalb von 1,5 Stunden nach vereinbarter Zeit abgeholt, besteht kein Anspruch mehr.'),
    ('\u00a75 Miete & Zahlung', 'Miete ist im Voraus für die gesamte Vertragslaufzeit zu zahlen. Alle Preise inkl. MwSt. Reinigungspauschale bei verschmutzter Rückgabe: 49,\u2013 \u20ac zzgl. MwSt.'),
    ('\u00a76 Gebrauch', 'Die Mietsache darf nur zum vereinbarten Zweck verwendet werden. Weitervermietung oder -verleihung ohne Erlaubnis ist verboten.'),
    ('\u00a77 Obliegenheiten', 'Pflegliche Behandlung, keine Entfernung von Kennzeichnungen, fachkundiger Auf- und Abbau gemäß Betriebsanleitung.'),
    ('\u00a78 Änderungen', 'Änderungen durch den Mieter bedürfen der vorherigen Zustimmung des Vermieters.'),
    ('\u00a79 Mängelrechte', 'Mängel sind unverzüglich anzuzeigen. Eine Kündigung ist erst nach erfolgloser Mängelbeseitigung zulässig.'),
    ('\u00a710 Haftung', 'Der Mieter übernimmt die Haftung für Sach- und Personenschäden. Eine entsprechende Versicherung ist abzuschließen.'),
    ('\u00a711 Vertragslaufzeit', 'Das Mietverhältnis wird befristet geschlossen und endet automatisch. Kündigung bedarf der Textform (E-Mail).'),
    ('\u00a712 Kaution', 'Je Mietsache ist eine Kaution von 150,\u2013 \u20ac bei Abholung zu hinterlegen. Rückerstattung bei sauberer Rückgabe mit Videobeweis.'),
    ('\u00a713 Rückgabe', 'Mietsache in ordnungsgemäßem Zustand zurückgeben. Bei Überschreitung der Mietdauer: tagesanteiliger Betrag.'),
]
for title, text in agb:
    p_a = doc.add_paragraph()
    spacing(p_a, 60, 20)
    add_run(p_a, title + ': ', True, 9.5, ORANGE)
    add_run(p_a, text, False, 9.5)

doc.add_page_break()

section_header(doc, 'STORNIERUNGSBEDINGUNGEN', 'CC0000')
t_st = doc.add_table(rows=5, cols=2)
full_width(t_st)
st_d = [
    ('Stornierungszeitpunkt', 'Kosten'),
    ('Bis 14 Tage vor dem Reservierungsdatum', 'Kostenlos'),
    ('7\u201314 Tage vorher', 'Pauschale 49,\u2013 \u20ac'),
    ('Weniger als 7 Tage vorher', 'Volle Mietkosten'),
    ('Nicht abgeholt ohne Stornierung', '1,5-facher Gesamtmietpreis'),
]
for i, (c1, c2) in enumerate(st_d):
    for c, txt in zip(t_st.rows[i].cells, [c1, c2]):
        c.text = txt
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        c.paragraphs[0].runs[0].font.name = 'Calibri'
        if i == 0:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            cell_bg(c, 'CC0000')
        else:
            c.paragraphs[0].runs[0].font.color.rgb = DARK
            cell_bg(c, 'FFF5F5' if i % 2 == 0 else 'FFFFFF')
    if i > 0:
        t_st.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        t_st.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RED

orange_line(doc)
section_header(doc, 'WEITERE BESTIMMUNGEN')
weitere = [
    ('\u00a715 Anwendbares Recht', 'Es gilt das Recht der Bundesrepublik Deutschland.'),
    ('\u00a716 Gerichtsstand', 'Gerichtsstand ist der Geschäftssitz des Vermieters (Vilshofen).'),
    ('\u00a717 Streitbeilegung', 'EU-Plattform: https://ec.europa.eu/consumers/odr \u2014 Der Vermieter ist bereit, an einem Schlichtungsverfahren teilzunehmen.'),
    ('\u00a718 Wetter', 'Bei Schlechtwetter kann der Vermieter stornieren. Kann das Material wegen Schlechtwetter nicht genutzt werden, besteht kein Entschädigungsanspruch.'),
]
for t, x in weitere:
    p_w = doc.add_paragraph()
    spacing(p_w, 60, 20)
    add_run(p_w, t + ': ', True, 9.5, ORANGE)
    add_run(p_w, x, False, 9.5)

doc.add_page_break()

# ===== VOLLMACHT =====
section_header(doc, 'EINMALIGE VOLLMACHT (nur wenn Abholer \u2260 Mieter)')
para(doc, 'Alle Angaben sind Pflichtangaben. Identität durch gültigen Lichtbildausweis nachweisen.', size=9, color=GRAY, b=60, a=60)

t_vm = doc.add_table(rows=1, cols=2)
full_width(t_vm)
for c, lbl in zip(t_vm.rows[0].cells, ['VOLLMACHTGEBER (Mieter)', 'ABHOLER (Bevollmächtigter)']):
    cell_bg(c, 'FF6B00')
    add_run(c.paragraphs[0], lbl, True, 10, WHITE)

t_vm2 = doc.add_table(rows=1, cols=2)
full_width(t_vm2)
for col_idx, labels in enumerate([
    ['Name, Vorname','Firma','Straße + Hausnummer','PLZ + Ort','Telefon / Mobil','E-Mail','Geburts-Datum'],
    ['Name, Vorname','Straße + Hausnummer','PLZ + Ort','Geburts-Datum']
]):
    c = t_vm2.rows[0].cells[col_idx]
    for label in labels:
        p_l = c.add_paragraph()
        spacing(p_l, 30, 2)
        add_run(p_l, label, True, 8, GRAY)
        p_f = c.add_paragraph()
        spacing(p_f, 0, 8)
        bottom_border_para(p_f)

for txt in [
    '\u2610  Hiermit bevollmächtige ich die nebenstehende Person, die für mich bestimmten Mietsachen in Empfang zu nehmen.',
    '\u2610  Hiermit bevollmächtige ich die nebenstehende Person, die für mich bestimmten Mietsachen an den Vermieter zu übergeben.',
]:
    p_v = doc.add_paragraph()
    spacing(p_v, 30, 20)
    add_run(p_v, txt, False, 9.5)

para(doc, 'Gültig einmalig in der Angelegenheit und für einen Tag.', size=9, color=GRAY, b=20, a=60)
orange_line(doc)
t_vm3 = doc.add_table(rows=2, cols=2)
no_borders(t_vm3)
full_width(t_vm3)
for c, lbl in zip(t_vm3.rows[0].cells, ['Event-Datum','Datum der Unterschrift']):
    add_run(c.paragraphs[0], lbl, True, 8, GRAY)
    c.add_paragraph('___ . ___ . _________').runs[0].font.size = Pt(10)
for c, lbl in zip(t_vm3.rows[1].cells, ['Unterschrift Vollmachtgeber / Mieter','Unterschrift Abholer']):
    top_border_para(c.paragraphs[0])
    add_run(c.paragraphs[0], lbl, True, 9.5, ORANGE)

doc.add_page_break()

# ===== ÜBERGABEPROTOKOLL =====
section_header(doc, 'ÜBERGABE- UND RÜCKGABEPROTOKOLL')
t_pr = doc.add_table(rows=2, cols=4)
full_width(t_pr)
for c, h in zip(t_pr.rows[0].cells, ['Mietgegenstand','\u2713 OK bei Abholung','Schäden','Verschmutzt']):
    cell_bg(c, 'FF6B00')
    add_run(c.paragraphs[0], h, True, 9.5, WHITE)
row_data = ['Hüpfburg komplett','\u2610','\u2610','\u2610']
for c, txt in zip(t_pr.rows[1].cells, row_data):
    c.text = txt
    c.paragraphs[0].runs[0].font.size = Pt(13 if '\u2610' in txt else 9.5)
    c.paragraphs[0].runs[0].font.name = 'Calibri'
    cell_bg(c, 'F9F9F9')
    if '\u2610' in txt:
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

para(doc, '', b=40, a=20)
t_dates = doc.add_table(rows=2, cols=2)
full_width(t_dates)
for c, lbl in zip(t_dates.rows[0].cells, ['ÜBERGABE-DATUM','RÜCKGABE-DATUM']):
    cell_bg(c, 'FFF3E0')
    add_run(c.paragraphs[0], lbl, True, 9, ORANGE)
for c in t_dates.rows[1].cells:
    c.text = '___ . ___ . _________'
    c.paragraphs[0].runs[0].font.size = Pt(11)
    p2 = c.add_paragraph('Pünktlich zurückgebracht:  \u2610 Ja   \u2610 Nein')
    p2.runs[0].font.size = Pt(9.5)
    p2.runs[0].font.name = 'Calibri'

para(doc, '', b=40, a=20)
para(doc, 'Für aufgeführte Schäden, fehlende Gegenstände oder Verschmutzungen wird folgendes in Rechnung gestellt:', True, b=60, a=20)
for _ in range(5):
    p_l = doc.add_paragraph()
    spacing(p_l, 20, 20)
    bottom_border_para(p_l, 'CCCCCC')

orange_line(doc)
t_final = doc.add_table(rows=2, cols=3)
no_borders(t_final)
full_width(t_final)
for c in t_final.rows[0].cells:
    c.add_paragraph()
for c, lbl in zip(t_final.rows[1].cells, ['Unterschrift Mieter','Unterschrift Abholer mit Vollmacht','Unterschrift Vermieter']):
    top_border_para(c.paragraphs[0])
    add_run(c.paragraphs[0], lbl, True, 9, ORANGE)

doc.save('c:/Users/Gutse/Projekte/Hupfgaudi/Mietvertrag_NEU.docx')
print('OK - Mietvertrag_NEU.docx gespeichert')
